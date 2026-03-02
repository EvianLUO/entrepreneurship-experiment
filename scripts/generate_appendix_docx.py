# -*- coding: utf-8 -*-
"""生成实验问卷题项附录 Word 文档"""
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请先安装: pip install python-docx")
    sys.exit(1)

OUTPUT_PATH = "/Users/eviane/Desktop/人工智能 创业/data/study 4/实验问卷题项附录.docx"

# 前测
PRETEST = {
    "title": "附录 实验问卷题项",
    "section1_title": "一、前测问卷",
    "section1_subtitle": "（一）基本信息与控制变量",
    "A1": ("A1. 您的创业次数", ["0次", "1次", "2次", "3次", "4次", "5次", "6次以上"]),
    "A2": ("A2. 您的总创业年限（仅填数字，单位：年）", None),
    "A3": ("A3. 您的创业阶段", ["准备期", "初创期", "成长期", "扩张期"]),
    "A4": ("A4. 您的企业成立时间", ["小于6个月", "6个月-1年", "1-2年", "2-3年", "3-5年", "5年以上"]),
    "A5": ("A5. 当前员工人数", ["0-5人", "6-20人", "21-50人", "50人以上"]),
    "A6": ("A6. 您目前从事的行业：（单选）", [
        "制造业", "建筑业", "物流交通", "教育/培训", "互联网", "计算机/软件", "批发和零售",
        "住宿餐饮业", "金融业", "房地产业", "出租/租赁", "专业服务（如法律/咨询服务）",
        "科学研究", "生活服务（如家政/美发类服务）", "医疗卫生/社会保障", "文化娱乐",
        "政府事业单位", "农、林、牧、渔业", "水利环境公共设施管理", "电力燃气", "采矿业",
        "国际组织", "其他（请注明）",
    ]),
    "A7": ("A7. 您从哪一年开始用生成式人工智能？（2022年11月30日ChatGPT发布）", ["2022年", "2023年", "2024年", "2025年"]),
    "A8": ("A8. 您平均每天使用生成式AI工具的频率", ["1-2次", "3-5次", "6-10次", "10次以上"]),
    "A9": ("A9. 您平均每次使用生成式AI的时长", ["5分钟以内", "5-15分钟", "15-30分钟", "30分钟-1小时", "1小时以上"]),
    "A10": ("A10. 本研究非常关注数据的质量。请证明您在认真作答，不要选择「苹果」，而是选择「香蕉」。", ["苹果", "橘子", "葡萄", "香蕉", "西瓜"]),
    "section2_subtitle": "（二）GenAI 依赖量表",
    "scale_7": "量表：1 = 非常不同意，7 = 非常同意",
    "genai_items": [
        "使用生成式AI时，我感到安心和放松",
        "如果生成式AI不可用，我会感到焦虑",
        "我对生成式AI有情感上的依赖",
        "生成式AI让我感到被支持",
        "我依赖生成式AI来帮助我思考问题",
        "没有生成式AI，我觉得自己的思维能力下降了",
        "我越来越依赖生成式AI来做决定",
        "我发现自己在很多任务中都依赖生成式AI",
        "我每天都会使用生成式AI",
        "我在工作和生活中频繁使用生成式AI",
        "我很难想象没有生成式AI的生活",
        "遇到问题时，我的第一反应是求助于生成式AI",
    ],
}

# 后测
POSTTEST = {
    "section_title": "二、后测问卷",
    "task_inv_title": "（一）任务投入度",
    "task_inv_intro": "请回想您刚才完成的决策任务。量表：1 = 非常不同意，7 = 非常同意",
    "task_inv_items": [
        "我认真对待了这个决策任务",
        "我在这个任务上投入了很多精力",
        "这个任务对我来说很重要",
        "我尽了最大努力来完成这个任务",
    ],
    "manip_title": "（二）操纵检验",
    "manip_m1": "M1. 在本次实验中，您完成决策任务的顺序是：（单选）",
    "manip_m1_opts": ["先独立撰写方案，然后咨询AI助手", "先咨询AI助手，然后独立撰写方案"],
    "manip_m2": "M2. 在做出最终决策之前，您在多大程度上先进行了独立思考？",
    "manip_m2_scale": "1 = 完全没有独立思考，直接依赖AI；7 = 完全独立思考后才参考AI",
    "manip_m3": "M3. AI助手在给出建议时，在多大程度上解释了它的推理过程和依据？",
    "manip_m3_scale": "1 = 完全没有解释，只给结论；7 = 非常详细地解释了原因和逻辑",
    "manip_m4": "M4. 您能理解AI为什么给出这样的建议吗？",
    "manip_m4_scale": "1 = 完全不理解；7 = 完全理解",
    "effect_title": "（三）效果逻辑量表",
    "effect_intro": "请回想您刚才完成的决策任务，以下每道题描述了两种不同的决策方式。请根据您在任务中的实际做法，选择更符合您的一端。1 = 左侧表述，7 = 右侧表述。",
    "effect_dim1": "维度一：手段导向 vs. 目标导向",
    "effect_dim2": "维度二：可承受损失 vs. 预期回报",
    "effect_dim3": "维度三：合作伙伴 vs. 竞争分析",
    "effect_dim4": "维度四：接纳意外 vs. 克服意外",
    "effect_items": [
        ("1", "我的决策是基于我们现有的手段/资源来定义项目的。", "我的决策是基于给定的项目目标来定义项目的。"),
        ("2", "在项目初期，我的目标定义是模糊的。", "在项目初期，我的目标定义是清晰的。"),
        ("3", "我以给定的手段/资源作为项目的起点。", "我以给定的项目目标作为项目的起点。"),
        ("4", "我围绕现有手段/资源逐步收敛至一个目标。", "我根据项目目标来确定所需的手段/资源。"),
        ("5", "我更多从现有手段出发，而非从明确的目标出发。", "我从一个清晰给定的项目目标出发。"),
        ("6", "我主要依据给定的资源来规划项目。", "我主要依据给定的目标来规划项目。"),
        ("7", "给定的手段深刻地影响了我的项目框架。", "给定的项目目标深刻地影响了我的项目框架。"),
        ("8", "我以潜在损失作为选择方案的决定性因素。", "我以潜在回报作为选择方案的决定性因素。"),
        ("9", "我基于可接受的损失来审批项目预算。", "我基于预期回报的计算来审批项目预算。"),
        ("10", "我主要考虑如何最小化风险和成本。", "我主要分析未来的回报。"),
        ("11", "我主要考量项目的潜在风险。", "我主要考量项目的潜在收益。"),
        ("12", "我基于潜在的损失风险来决策资本支出。", "我基于潜在的回报来决策资本支出。"),
        ("13", "我通过内/外部伙伴关系及协议来降低项目风险。", "我通过彻底的市场和竞争者分析来识别项目风险。"),
        ("14", "我与伙伴/利益相关者基于各自能力共同决策。", "我基于系统的市场分析来做出决策。"),
        ("15", "我侧重通过接触潜在伙伴和客户来降低风险。", "我侧重通过市场分析来早期识别风险，以便调整策略。"),
        ("16", "为降低风险，我积极开展合作并寻求预先承诺。", "为识别风险，我专注于市场分析和预测。"),
        ("17", "我会尝试在过程中整合意外的结果和发现，即使与初始目标不完全一致。", "我只在初始目标面临风险时，才整合意外的发现。"),
        ("18", "我的项目过程足够灵活，可以根据新发现调整。", "我的项目过程专注于毫无延迟地达成目标。"),
        ("19", "新的研究发现会影响我的项目目标。", "新的研究发现不会影响我的项目目标。"),
        ("20", "我的项目规划是在执行过程中以小步快跑的方式进行的。", "我的项目规划基本上在项目开始时就已制定完成。"),
        ("21", "即使可能导致延迟，我仍保持灵活并利用出现的机会。", "我首先确保毫无延迟地达成最初设定的目标。"),
        ("22", "我允许项目随着机会的出现而演变，即使这些机会与初始目标不符。", "我始终关注于达成最初的项目目标。"),
    ],
    "basic_title": "（四）基本信息",
    "basic_gender": "您的性别为",
    "basic_gender_opts": ["男", "女"],
    "basic_age": "您的年龄是",
    "basic_age_opts": ["18-25岁", "26-30岁", "31-35岁", "36-40岁", "41-45岁", "46-50岁", "51-55岁", "56-60岁", "60岁以上"],
    "basic_edu": "请选择您的最高学历",
    "basic_edu_opts": ["小学及以下", "初中", "普高/中专/技校/职高", "专科", "本科", "硕士", "博士"],
}


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.size = Pt(14) if level == 1 else Pt(12)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    if bold:
        run.bold = True
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)

    add_heading(doc, PRETEST["title"], level=0)
    doc.add_paragraph()

    # ========== 前测 ==========
    add_heading(doc, PRETEST["section1_title"], level=1)
    add_heading(doc, PRETEST["section1_subtitle"], level=2)

    for key in ["A1", "A2", "A3", "A4", "A5", "A6", "A7", "A8", "A9", "A10"]:
        q, opts = PRETEST[key]
        add_para(doc, q, bold=True)
        if opts:
            add_para(doc, "选项：" + "；".join(opts))

    add_heading(doc, PRETEST["section2_subtitle"], level=2)
    add_para(doc, PRETEST["scale_7"])
    for i, item in enumerate(PRETEST["genai_items"], 1):
        add_para(doc, f"G{i}. {item}")

    doc.add_paragraph()

    # ========== 后测 ==========
    add_heading(doc, POSTTEST["section_title"], level=1)

    add_heading(doc, POSTTEST["task_inv_title"], level=2)
    add_para(doc, POSTTEST["task_inv_intro"])
    for i, item in enumerate(POSTTEST["task_inv_items"], 1):
        add_para(doc, f"{i}. {item}")

    add_heading(doc, POSTTEST["manip_title"], level=2)
    add_para(doc, POSTTEST["manip_m1"])
    add_para(doc, "选项：" + "；".join(POSTTEST["manip_m1_opts"]))
    add_para(doc, POSTTEST["manip_m2"])
    add_para(doc, POSTTEST["manip_m2_scale"])
    add_para(doc, POSTTEST["manip_m3"])
    add_para(doc, POSTTEST["manip_m3_scale"])
    add_para(doc, POSTTEST["manip_m4"])
    add_para(doc, POSTTEST["manip_m4_scale"])

    add_heading(doc, POSTTEST["effect_title"], level=2)
    add_para(doc, POSTTEST["effect_intro"])
    add_para(doc, POSTTEST["effect_dim1"])
    for num, left, right in POSTTEST["effect_items"][:7]:
        add_para(doc, f"B{num}. 1：{left}；7：{right}")
    add_para(doc, POSTTEST["effect_dim2"])
    for num, left, right in POSTTEST["effect_items"][7:12]:
        add_para(doc, f"B{num}. 1：{left}；7：{right}")
    add_para(doc, POSTTEST["effect_dim3"])
    for num, left, right in POSTTEST["effect_items"][12:16]:
        add_para(doc, f"B{num}. 1：{left}；7：{right}")
    add_para(doc, POSTTEST["effect_dim4"])
    for num, left, right in POSTTEST["effect_items"][16:22]:
        add_para(doc, f"B{num}. 1：{left}；7：{right}")

    add_heading(doc, POSTTEST["basic_title"], level=2)
    add_para(doc, POSTTEST["basic_gender"])
    add_para(doc, "选项：" + "；".join(POSTTEST["basic_gender_opts"]))
    add_para(doc, POSTTEST["basic_age"])
    add_para(doc, "选项：" + "；".join(POSTTEST["basic_age_opts"]))
    add_para(doc, POSTTEST["basic_edu"])
    add_para(doc, "选项：" + "；".join(POSTTEST["basic_edu_opts"]))

    doc.save(OUTPUT_PATH)
    print("已生成：" + OUTPUT_PATH)


if __name__ == "__main__":
    main()
